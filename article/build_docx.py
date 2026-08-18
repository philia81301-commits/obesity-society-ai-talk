"""把 content/article-newsletter-2026.md 排成投稿用的 DOCX。

版面完全比照學會既有的醫學專論體例（自 2025 年〈體重管理的數位轉型〉一文實測）：
A4、上下 2.54cm／左右 3.17cm、內文 12pt、中文新細明體、英文 Calibri，
標題置中、作者與單位右對齊，內文引註為上標數字，文末 Vancouver 格式參考資料。

用法：python build_docx.py [輸出路徑.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD = ROOT.parent / "content" / "article-newsletter-2026.md"
FIGS = ROOT / "figs"
IMAGES = ROOT.parent / "slides" / "images"

CN_FONT = "新細明體"
EN_FONT = "Calibri"
BODY_PT = 12
TEXT_WIDTH_CM = 21.0 - 3.17 * 2  # 14.66

# 圖片：編號 → (檔案, 版面寬度 cm)
FIGURES = {
    "圖一": (FIGS / "fig1_flow.png", TEXT_WIDTH_CM),
    # 圖二 2026-08-18 由實際操作截圖改為概念圖，避免欄位配置與貼稿樣板被照抄
    "圖二": (FIGS / "fig2_datalayer.png", TEXT_WIDTH_CM),
    "圖三": (IMAGES / "p25_report_charts.png", TEXT_WIDTH_CM),
    "圖四": (FIGS / "fig4_clinic_steps.png", TEXT_WIDTH_CM),
}

# 圖表插入點：在「包含這段文字的段落」之後插入。
# ⚠️ 錨點要挑不會因潤稿而變動的短句；改內文後若錨點失配，
#    build() 結尾的檢查會直接丟出例外，不會默默少一張圖。
ANCHORS = [
    ("兩者並列即可看出定位差異", "表一"),
    ("說明導入的實際樣貌", "圖一"),
    ("而在做出來的成本太高", "表二"),
    ("自需求提出到上線", "圖二"),
    ("分析輸出一律為聚合統計", "圖三"),
    ("與病人一起決定下一步", "圖四"),
]


# ---------- 基礎樣式 ----------

def set_run(run, *, size=BODY_PT, bold=False, italic=False, sup=False, color=None):
    run.font.size = Pt(size)
    run.font.name = EN_FONT
    run.font.bold = bold
    run.font.italic = italic
    run.font.superscript = sup
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


TOKEN = re.compile(r"(\*\*.+?\*\*|\^\[\d+\]|\*[^*]+?\*)")


def add_rich(par, text, *, size=BODY_PT, base_bold=False):
    """處理 **粗體**、*斜體*、^[n] 上標引註。"""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run(par.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("^["):
            set_run(par.add_run(part[2:-1]), size=size, sup=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            set_run(par.add_run(part[1:-1]), size=size, italic=True)
        else:
            set_run(par.add_run(part), size=size, bold=base_bold)


def para(doc, text="", *, align=None, indent=True, space_after=6,
         size=BODY_PT, bold=False, line=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    if text:
        add_rich(p, text, size=size, base_bold=bold)
    return p


# ---------- Markdown 解析 ----------

def load_blocks():
    raw = MD.read_text(encoding="utf8")
    parts = raw.split("\n---\n")
    body = parts[0]
    figsec = parts[1] if len(parts) > 1 else ""
    tail = parts[2] if len(parts) > 2 else ""
    return body, figsec, tail


def parse_tables(figsec):
    """把 `## 圖表` 區塊裡的每個 **表X …** 段落拆成 (標題, 表格, 註)。"""
    out = {}
    chunks = re.split(r"\n(?=\*\*表[一二三四])", figsec)
    for ch in chunks:
        m = re.match(r"\*\*(表([一二三四]).+?)\*\*", ch)
        if not m:
            continue
        rows = [ln for ln in ch.splitlines() if ln.strip().startswith("|")]
        rows = [r for r in rows if not re.match(r"^\|[\s|:-]+\|$", r.strip())]
        grid = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        note = re.search(r"^> (.+)$", ch, re.M)
        out["表" + m.group(2)] = (m.group(1), grid, note.group(1) if note else "")
    return out


def parse_fig_captions(figsec):
    caps = {}
    for m in re.finditer(r"\*\*(圖[一二三四])\*\*　?(.+)", figsec):
        caps[m.group(1)] = f"{m.group(1)}　{m.group(2).strip()}"
    return caps


# ---------- 插入圖表 ----------

def insert_figure(doc, key, caption):
    path, width = FIGURES[key]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True   # 圖與圖說不分頁
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = 1.2
    add_rich(cap, caption, size=10.5)


def insert_table(doc, caption, grid, source):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True   # 標題不與表格分家
    add_rich(cap, caption, size=10.5, base_bold=True)

    t = doc.add_table(rows=len(grid), cols=len(grid[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(grid):
        # 單一儲存格不跨頁；標題列跨頁時重複
        trPr = t.rows[i]._tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
        if i == 0:
            trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_rich(p, cell, size=10, base_bold=(i == 0))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.line_spacing = 1.2
    add_rich(note, source, size=9.5)


# ---------- 主流程 ----------

def build(dest: Path):
    body, figsec, tail = load_blocks()
    tables = parse_tables(figsec)
    fig_caps = parse_fig_captions(figsec)

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = EN_FONT
    st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(3.17)
    s.top_margin = s.bottom_margin = Cm(2.54)

    lines = body.strip().splitlines()
    title = lines[0].lstrip("# ").strip()

    para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
         size=16, bold=True, space_after=10)
    para(doc, "潘湘如　醫師", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False,
         size=12, space_after=0)
    para(doc, "高雄榮民總醫院家庭醫學部", align=WD_ALIGN_PARAGRAPH.RIGHT,
         indent=False, size=12, space_after=14)

    placed = set()

    def flush(text):
        """輸出一個段落，並在命中錨點時補上圖表。"""
        if text.startswith("## "):
            head = text[3:].strip()
            para(doc, head, indent=False, size=13, bold=True,
                 space_after=6, line=1.4)
            return
        para(doc, text)
        for anchor, key in ANCHORS:
            if anchor in text:
                placed.add(key)
                if key.startswith("表"):
                    insert_table(doc, *tables[key])
                else:
                    insert_figure(doc, key, fig_caps[key])

    blocks = [b.strip() for b in "\n".join(lines[1:]).split("\n\n") if b.strip()]
    # 第一塊是 md 裡的作者／單位兩行，上面已單獨排版，這裡略過
    if blocks and blocks[0].startswith("潘湘如"):
        blocks = blocks[1:]
    for blk in blocks:
        flush(blk.replace("\n", ""))

    # 附錄與參考資料
    for blk in [b.strip() for b in tail.split("\n\n") if b.strip()]:
        blk = blk.replace("\n", "")
        if blk.startswith("## "):
            para(doc, blk[3:].strip(), indent=False, size=13, bold=True,
                 space_after=6, line=1.4)
        elif re.match(r"^\d+\.（[A-D]）", blk):
            # 解答與解析：懸掛縮排，讓「1.（C）」凸出於解析文字之外
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Pt(30)
            pf.first_line_indent = Pt(-30)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.4
            add_rich(p, blk, size=11)
        elif re.match(r"^（[A-D]）", blk):
            # 自我評量的選項：不首行縮排，改用整段縮排讓題幹與選項分層
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Pt(24)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.4
            add_rich(p, blk, size=11.5)
        elif re.match(r"^\[\d+\]", blk):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Pt(24)
            pf.first_line_indent = Pt(-24)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.3
            add_rich(p, blk, size=10.5)
        else:
            para(doc, blk, size=11, line=1.4, space_after=4)

    missing = [k for _, k in ANCHORS if k not in placed]
    if missing:
        raise SystemExit(
            "錨點失配，下列圖表沒有被插入：" + "、".join(missing)
            + "（通常是潤稿時改到 ANCHORS 裡的句子，請更新錨點文字）")

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    print("已輸出：", dest)


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "【醫學專論】人工智慧在肥胖治療運用的新發展___潘湘如.docx"
    build(out)
